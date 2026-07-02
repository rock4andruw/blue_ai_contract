"""Contract parser: extracts structured clauses from MD, DOCX, and PDF files."""

import hashlib
import re
from pathlib import Path
from typing import Any, Dict, List, Optional
from dataclasses import dataclass, field


@dataclass
class ClauseElement:
    clause_number: Optional[str]
    title: str
    content: str
    page_number: int
    content_hash: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ContractDocument:
    filename: str
    file_type: str
    raw_text: str
    clauses: List[ClauseElement]
    metadata: Dict[str, Any] = field(default_factory=dict)


class ContractParser:
    # Matches: 第5條、第5.2條、5.、5.2、一、二、三、Article 3、Section 3.1
    CLAUSE_PATTERNS = [
        r'^第\s*(\d+(?:\.\d+)*)\s*條',
        r'^([一二三四五六七八九十百]+)[、．]',  # 一、二、三、（中文數字條款）
        r'^(\d+(?:\.\d+)+)\s*[、\s]',          # 5.2、或 5.2 開頭
        r'^(\d+)\.\s+\S',                       # 5. 開頭（整條）
        r'^Article\s+(\d+(?:\.\d+)*)',
        r'^Section\s+(\d+(?:\.\d+)*)',
    ]

    def parse_file(self, file_path: str) -> ContractDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".md":
            return self._parse_md(file_path)
        elif suffix == ".pdf":
            return self._parse_pdf(file_path)
        elif suffix == ".docx":
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported format: {suffix}")

    # ------------------------------------------------------------------
    # MD parser (primary for demo)
    # ------------------------------------------------------------------

    @staticmethod
    def _strip_track_changes(text: str) -> str:
        """Remove Word/HTML track-change markup from text.

        <del ...>...</del>  → removed (deleted content)
        <ins ...>...</ins>  → inner text kept (accepted insertion)
        Any remaining HTML tags → stripped
        """
        text = re.sub(r'<del[^>]*>.*?</del>', '', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<ins[^>]*>(.*?)</ins>', r'\1', text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<[^>]+>', '', text)
        return text

    def _parse_md(self, file_path: str) -> ContractDocument:
        raw = Path(file_path).read_text(encoding="utf-8")
        text = self._strip_track_changes(raw)
        clauses = self._split_md_clauses(text)
        return ContractDocument(
            filename=Path(file_path).name,
            file_type="md",
            raw_text=text,
            clauses=clauses,
            metadata={"clause_count": len(clauses)},
        )

    def _split_md_clauses(self, text: str) -> List[ClauseElement]:
        """Split MD into clauses by ## headings then by numbered sub-items."""
        clauses: List[ClauseElement] = []
        # Split on ## section headings
        sections = re.split(r'\n(?=## )', text)

        for section in sections:
            lines = section.strip().splitlines()
            if not lines:
                continue

            first_line = lines[0].strip()
            # The first "section" (before the first "## ") may be the
            # document's H1 title + preamble, not a real clause. The
            # title/version label (e.g. "v1" vs "v4") isn't a negotiable
            # contract term — diffing it as a clause creates a false
            # modified-clause diff whenever only the title changes, even if
            # the actual preamble body (甲方/乙方/開場白) is identical.
            # Compare the preamble body only, or skip if there is none.
            if first_line.startswith('#') and not first_line.startswith('##'):
                preamble_body = "\n".join(lines[1:]).strip()
                if not preamble_body:
                    continue
                clauses.append(ClauseElement(
                    clause_number=None,
                    title="文件序言",
                    content=preamble_body,
                    page_number=1,
                    content_hash=self._md5(preamble_body),
                ))
                continue

            # Section heading becomes a parent clause
            heading = lines[0].lstrip("#").strip()
            section_number = self._extract_clause_number(heading)
            section_body = "\n".join(lines[1:]).strip()

            # Try to split body into sub-clauses by numbered lines (e.g. "4.1 ...")
            sub_clauses = self._split_numbered_items(section_body, parent_title=heading)

            if sub_clauses:
                # Add section header as its own entry (no content, just structure)
                clauses.append(ClauseElement(
                    clause_number=section_number,
                    title=heading,
                    content=heading,
                    page_number=1,
                    content_hash=self._md5(heading),
                ))
                clauses.extend(sub_clauses)
            else:
                # Reconstruct content from heading + body (both already stripped
                # of "## " markdown syntax) rather than the raw section text —
                # otherwise a section that gains/loses sub-clauses between two
                # document versions flips between "## " being present or absent
                # in its content, producing a false diff on an unchanged heading.
                full_content = heading + (f"\n{section_body}" if section_body else "")
                clauses.append(ClauseElement(
                    clause_number=section_number,
                    title=heading,
                    content=full_content,
                    page_number=1,
                    content_hash=self._md5(full_content),
                ))

        return clauses

    def _split_numbered_items(self, text: str, parent_title: str) -> List[ClauseElement]:
        """Split text block into sub-clauses by lines starting with N.N pattern."""
        pattern = re.compile(r'^(\d+\.\d+)\s+(.+)', re.MULTILINE)
        matches = list(pattern.finditer(text))
        if not matches:
            return []

        items: List[ClauseElement] = []
        for i, m in enumerate(matches):
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            block = text[start:end].strip()
            number = m.group(1)
            title_text = m.group(2).strip()[:60]
            items.append(ClauseElement(
                clause_number=number,
                title=f"{number} {title_text}",
                content=block,
                page_number=1,
                content_hash=self._md5(block),
            ))
        return items

    # ------------------------------------------------------------------
    # PDF parser
    # ------------------------------------------------------------------

    def _parse_pdf(self, file_path: str) -> ContractDocument:
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber not installed: pip install pdfplumber")

        chunks: List[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                chunks.append(page.extract_text() or "")

        raw = "\n".join(chunks)
        clauses = self.split_into_clauses(raw)
        return ContractDocument(
            filename=Path(file_path).name,
            file_type="pdf",
            raw_text=raw,
            clauses=clauses,
            metadata={"page_count": len(chunks)},
        )

    # ------------------------------------------------------------------
    # DOCX parser
    # ------------------------------------------------------------------

    def _parse_docx(self, file_path: str) -> ContractDocument:
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx not installed: pip install python-docx")

        doc = docx.Document(file_path)

        # python-docx exposes doc.paragraphs and doc.tables as separate
        # collections — reading only paragraphs silently drops every native
        # Word table (payment schedules, pricing tiers, SLA metric tables),
        # which then never enters diffs at all, not even as a low-confidence
        # match. Walk the document body in original order so tables are
        # captured and interleaved with the paragraphs around them.
        parts = []
        for block in self._iter_docx_block_items(doc):
            if isinstance(block, docx.text.paragraph.Paragraph):
                if block.text.strip():
                    parts.append(block.text)
            else:
                table_text = self._docx_table_to_text(block)
                if table_text:
                    parts.append(table_text)
        raw = "\n\n".join(parts)
        clauses = self.split_into_clauses(raw)
        return ContractDocument(
            filename=Path(file_path).name,
            file_type="docx",
            raw_text=raw,
            clauses=clauses,
            metadata={},
        )

    @staticmethod
    def _iter_docx_block_items(doc):
        """Yield paragraphs and tables in original document order.

        python-docx's doc.paragraphs / doc.tables are separate flat lists
        with no relative ordering — walking the underlying XML body instead
        keeps a table interleaved with the paragraphs immediately around it
        (e.g. "付款排程如下表：" followed by the actual table).
        """
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        body = doc.element.body
        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                yield Paragraph(child, doc)
            elif child.tag == qn('w:tbl'):
                yield Table(child, doc)

    @staticmethod
    def _docx_table_to_text(table) -> str:
        """Render a Word table as plain text, one row per line, cells
        joined with ' | ', so table content flows through the same clause
        splitter as regular paragraphs instead of being silently dropped."""
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                rows.append(' | '.join(cells))
        return '\n'.join(rows)

    # ------------------------------------------------------------------
    # Shared utilities
    # ------------------------------------------------------------------

    def split_into_clauses(self, text: str, page_number: int = 1) -> List[ClauseElement]:
        """Generic clause splitter for plain text (PDF/DOCX fallback)."""
        clauses: List[ClauseElement] = []
        for para in text.split("\n\n"):
            para = para.strip()
            if not para:
                continue
            clauses.append(ClauseElement(
                clause_number=self._extract_clause_number(para),
                title=self._extract_title(para),
                content=para,
                page_number=page_number,
                content_hash=self._md5(para),
            ))
        return clauses

    def _extract_clause_number(self, text: str) -> Optional[str]:
        first_line = text.strip().splitlines()[0] if text.strip() else ""
        for pattern in self.CLAUSE_PATTERNS:
            m = re.match(pattern, first_line.strip())
            if m:
                return m.group(1)
        return None

    def _extract_title(self, text: str) -> str:
        lines = text.strip().splitlines()
        first = lines[0].strip() if lines else ""
        return first[:60] if len(first) > 5 else "未命名條款"

    @staticmethod
    def _md5(text: str) -> str:
        normalized = re.sub(r"\s+", "", text).lower()
        return hashlib.md5(normalized.encode("utf-8")).hexdigest()
